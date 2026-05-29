"""Raw-text extraction per file type, with an OCR fallback for scanned PDFs.

Each extractor returns a list of `Segment(page, text)`. PDFs yield one
segment per page (1-indexed) so citations can point at a page; Word and
Excel have no pages, so `page` is None.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

from services.errors import UnsupportedFileType

logger = logging.getLogger(__name__)


@dataclass
class Segment:
    page: int | None
    text: str


# ── PDF ────────────────────────────────────────────────────────────
def _ocr_page(page) -> str:
    """Best-effort OCR for a scanned PDF page. Returns '' if OCR is
    unavailable (no system tesseract) so ingestion never hard-fails."""
    try:
        import pytesseract

        image = page.to_image(resolution=200).original
        return pytesseract.image_to_string(image).strip()
    except Exception as exc:  # tesseract missing, render failure, etc.
        logger.warning("OCR fallback failed on page %s: %s", page.page_number, exc)
        return ""


def extract_pdf(path: Path) -> list[Segment]:
    import pdfplumber

    segments: list[Segment] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:  # likely a scanned/image page → OCR
                text = _ocr_page(page)
            if text:
                segments.append(Segment(page=i, text=text))
    return segments


# ── Word ───────────────────────────────────────────────────────────
def extract_docx(path: Path) -> list[Segment]:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # include table cell text, which python-docx keeps separate
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    return [Segment(page=None, text=text)] if text else []


# ── Excel ──────────────────────────────────────────────────────────
def extract_xlsx(path: Path) -> list[Segment]:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    segments: list[Segment] = []
    for sheet in wb.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            # prefix the sheet name so multi-sheet context survives
            segments.append(Segment(page=None, text=f"[{sheet.title}]\n" + "\n".join(rows)))
    wb.close()
    return segments


_EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
}


def extract_segments(path: Path) -> list[Segment]:
    ext = path.suffix.lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise UnsupportedFileType(f"No extractor for '{ext}' files")
    return extractor(path)
